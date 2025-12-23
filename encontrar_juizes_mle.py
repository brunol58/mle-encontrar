import streamlit as st
import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import os

st.set_page_config(page_title="Relatórios de MLEs", layout="wide")

st.title("Extração de Juízes e Geração de Relatórios Word")
st.info("Aplicação em fase de testes - Desenvolvido por Bruno Ferreira da Silva")

# =========================
# Upload do CSV
# =========================
uploaded_file = st.file_uploader("Envie a planilha CSV", type=["csv"])

if uploaded_file is not None:
    if "df" not in st.session_state:
        df = pd.read_csv(uploaded_file, sep=';', encoding='utf-8', dtype={'Número do Processo': str})
        df["Número do Processo"] = df["Número do Processo"].str.strip("\t")
        df["Número do Mandado"] = df["Número do Mandado"].str.strip("\t")
        df['Número do Processo Mod'] = df['Número do Processo'].str.replace('826', '', regex=False)
        st.session_state.df = df
    else:
        df = st.session_state.df

    st.subheader("Preview da Planilha")
    st.dataframe(df.head())

    BASE_URL = "https://esaj.tjsp.jus.br"

    # =========================
    # Funções de extração
    # =========================
    def formatar_numero_cnj(numero):
        return f"{numero[:7]}-{numero[7:9]}.{numero[9:13]}.8.26.{numero[13:]}"

    def gerar_link(numero_mod):
        numero_formatado = formatar_numero_cnj(numero_mod)
        foro = numero_mod[-4:]
        return (
            f"https://esaj.tjsp.jus.br/cpopg/search.do?"
            f"conversationId=&cbPesquisa=NUMPROC"
            f"&numeroDigitoAnoUnificado={numero_mod}"
            f"&foroNumeroUnificado={foro}"
            f"&dadosConsulta.valorConsultaNuUnificado={numero_formatado}"
            f"&dadosConsulta.valorConsultaNuUnificado=UNIFICADO"
            f"&dadosConsulta.valorConsulta="
            f"&dadosConsulta.tipoNuProcesso=UNIFICADO"
        )

    def extrair_juiz(numero_mod):
        headers = {"User-Agent": "Mozilla/5.0"}
        try:
            resp = requests.get(gerar_link(numero_mod), headers=headers, timeout=10)
            if resp.status_code != 200:
                return f"Erro HTTP {resp.status_code}"
            soup = BeautifulSoup(resp.text, "html.parser")
            proc_princ = soup.find("a", class_="processoPrinc")
            if proc_princ:
                href_princ = proc_princ.get("href")
                if not href_princ:
                    return "Link do processo principal não encontrado"
                resp2 = requests.get(BASE_URL + href_princ, headers=headers, timeout=10)
                soup2 = BeautifulSoup(resp2.text, "html.parser")
                juiz_princ = soup2.find("span", id="juizProcesso")
                return juiz_princ.get_text(strip=True) if juiz_princ else "Juiz não encontrado"
            else:
                juiz = soup.find("span", id="juizProcesso")
                return juiz.get_text(strip=True) if juiz else "Juiz não encontrado"
        except:
            return "Erro ou não encontrado"

    # =========================
    # Botão para extrair juízes
    # =========================
    if st.button("Extrair juízes"):
        with st.spinner("Extraindo juízes..."):
            resultados_juiz = []
            for processo in df["Número do Processo Mod"]:
                juiz = extrair_juiz(processo)
                resultados_juiz.append(juiz)
                time.sleep(1)  # evitar bloqueio
            df["Juiz"] = resultados_juiz
            st.session_state.df = df
        st.success("Extração concluída!")
        st.dataframe(df[["Número do Processo", "Juiz"]])

    # =========================
    # Preenchimento manual
    # =========================
    if "Juiz" in df.columns:
        juizes_nao_encontrados = df[df["Juiz"] == "Juiz não encontrado"]
        if not juizes_nao_encontrados.empty:
            st.warning("Alguns juízes não foram encontrados. Complete manualmente:")
            for i, row in juizes_nao_encontrados.iterrows():
                juiz_manual = st.text_input(f"Número do processo: {row['Número do Processo']}", "")
                if juiz_manual:
                    df.at[i, "Juiz"] = juiz_manual
            st.session_state.df = df

    # =========================
    # Geração de relatórios Word
    # =========================
    if st.button("Gerar relatórios Word"):
        if "Juiz" not in df.columns:
            st.error("Extraia os juízes antes de gerar os relatórios.")
        else:
            os.makedirs("relatorios_juizes_word", exist_ok=True)
            for juiz, grupo in df.groupby("Juiz"):
                if juiz == "Erro ou não encontrado":
                    continue
                word_filename = f"relatorios_juizes_word/{juiz.replace('/', '_').replace(' ', '_')}.docx"
                doc = Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(12)

                title = doc.add_paragraph(f"MLEs para assinatura - {juiz}")
                title.style = doc.styles['Heading 1']
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph()

                for vara, processos in grupo.sort_values("Órgão/Vara").groupby("Órgão/Vara"):
                    subtitle = doc.add_paragraph(f"Vara: {vara}")
                    subtitle.style = doc.styles['Heading 2']
                    for _, row in processos.iterrows():
                        doc.add_paragraph(row['Número do Processo'].strip())
                    doc.add_paragraph()
                doc.save(word_filename)
            st.success("Relatórios Word gerados em relatorios_juizes_word/")
            st.info("Os arquivos estão disponíveis na pasta relatorios_juizes_word no ambiente do Streamlit Cloud.")
